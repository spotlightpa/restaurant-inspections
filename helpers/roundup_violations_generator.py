import os
import re
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


AP_MONTHS = {
    "January": "Jan.", "February": "Feb.", "March": "March",
    "April": "April", "May": "May", "June": "June",
    "July": "July", "August": "Aug.", "September": "Sept.",
    "October": "Oct.", "November": "Nov.", "December": "Dec."
}

AP_MONTHS_REVERSE = {v: k for k, v in AP_MONTHS.items()}

AP_NUMS = {
    1: "One", 2: "Two", 3: "Three", 4: "Four", 5: "Five",
    6: "Six", 7: "Seven", 8: "Eight", 9: "Nine"
}


def reverse_ap_date(s):
    if not isinstance(s, str):
        return pd.NaT
    for abbr, full in AP_MONTHS_REVERSE.items():
        s = s.replace(abbr, full)
    return pd.to_datetime(s, errors="coerce")


def ap_date(dt, include_month=True):
    month = AP_MONTHS[dt.strftime("%B")]
    day = str(dt.day)
    return f"{month} {day}" if include_month else day


def get_week_range():
    today = datetime.today()
    start = today - timedelta(days=today.weekday() + 7)  # Monday of last week
    end = start + timedelta(days=6)  # Sunday of last week
    if start.month == end.month:
        date_range = f"{ap_date(start)}-{end.day}"
    else:
        date_range = f"{ap_date(start)}-{ap_date(end)}"
    return date_range, start.strftime("%Y-%m-%d"), start, end


def generate_roundup_from_violations(roundup_path, county_slug):
    try:
        df = pd.read_excel(roundup_path)

        # Join ai_summary + risk_level from main inspections file
        inspections_path = "data/inspections.xlsx"
        try:
            insp = pd.read_excel(inspections_path, dtype=str)
            join_cols = [c for c in ["facility", "inspection_date", "ai_summary", "risk_level"] if c in insp.columns]
            if "facility" in join_cols and "inspection_date" in join_cols:
                insp = insp[join_cols].drop_duplicates(subset=["facility", "inspection_date"])
                df = df.merge(insp, on=["facility", "inspection_date"], how="left")
                print("Joined ai_summary and risk_level from inspections.xlsx")
                matched = df["ai_summary"].notna().sum()
                print(f"  ai_summary matched on {matched} of {len(df)} rows")
                if matched == 0:
                    print(f"  Sample roundup dates: {df['inspection_date'].head(3).tolist()}")
                    print(f"  Sample inspections dates: {insp['inspection_date'].head(3).tolist()}")
        except Exception as e:
            print(f"⚠️ Could not join from inspections.xlsx: {e}")

        # Filter to this county
        df = df[df["county"].str.strip().str.lower() == county_slug.lower()].copy()
        if df.empty:
            print(f"⚠️ No data for county: {county_slug}")
            return

        county_name = county_slug.title()
        date_range, date_slug, week_start, week_end = get_week_range()

        # Parse inspection dates
        df["_parsed_date"] = df["inspection_date"].apply(reverse_ap_date)

        # Filter to last completed week
        df = df[
            (df["_parsed_date"] >= pd.Timestamp(week_start.date())) &
            (df["_parsed_date"] <= pd.Timestamp(week_end.date()))
        ]

        if df.empty:
            print(f"⚠️ No inspections in date range for {county_slug}")
            return

        has_violations = df[df["comment"].fillna("").str.strip() != ""].copy()
        no_violations = df[df["comment"].fillna("").str.strip() == ""].copy()

        passed_deduped = no_violations.drop_duplicates(subset=["facility", "inspection_date"])

        # Build document
        doc = Document()

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title.add_run(f"{county_name} County Restaurant Inspections, {date_range}")
        run.bold = True
        run.font.size = Pt(20)

        # Date
        doc.add_paragraph(datetime.today().strftime("%B %-d, %Y"))

        # Intro
        intro = doc.add_paragraph(
            f"The state inspected the following food establishments in {county_name} County. "
            f"To find more restaurant safety information, view full inspection reports, and sign up for real-time text or email alerts "
            f"for specific restaurants or locations, visit Spotlight PA\u2019s "
        )
        add_hyperlink(intro, "Restaurant Safety Tracker", "https://www.spotlightpa.org/restaurant-inspections/")
        intro.add_run(".")

        intro2 = doc.add_paragraph("The ")
        add_hyperlink(intro2, "Pennsylvania Department of Agriculture", "https://www.pa.gov/agencies/pda/food/food-safety/retail-food-inspection-reports")
        intro2.add_run(" produces retail food inspection reports for 61 of Pennsylvania's 67 counties "
            f"using the FDA Model Food Code. Inspections are conducted regularly throughout the state, and results are posted as "
            f"inspections are conducted. As noted by the Department of Agriculture, inspections are a \u201csnapshot\u201d of a particular day. "
            f"Many violations are relatively minor and are fixed at the time of inspection."
        )

        # Out-of-compliance section
        heading_out = doc.add_paragraph()
        run_out = heading_out.add_run("Facilities with violations:")
        run_out.bold = True
        run_out.font.size = Pt(16)

        if has_violations.empty:
            doc.add_paragraph("No violations recorded.")
        else:
            for (facility, inspection_date), group in has_violations.groupby(["facility", "inspection_date"], sort=False):
                address = str(group.iloc[0].get("address", ""))
                date = str(inspection_date)

                facility_slug = re.sub(r'[^a-z0-9]+', '-', f"{facility} {address}".lower()).strip('-')
                tracker_url = f"https://www.spotlightpa.org/restaurant-inspections/#{facility_slug}"

                p = doc.add_paragraph()
                add_hyperlink(p, facility, tracker_url)
                p.add_run(f"\n{address}").italic = True
                p.add_run(f"\n{date}").italic = True

                # Count violations by risk level and show summary line
                if "risk_level" in group.columns:
                    all_levels = (
                        group["risk_level"]
                        .dropna()
                        .astype(str)
                        .str.split(r"\s*\|\s*")
                        .explode()
                        .str.strip()
                        .str.title()
                    )
                    all_levels = all_levels[~all_levels.isin(["Na", "Nan", ""])]
                    counts = all_levels.value_counts()
                    parts = []
                    for level in ["High Risk", "Moderate Risk", "Low Risk"]:
                        count = counts.get(level, 0)
                        if count:
                            clean_level = re.sub(r'\s*risk\s*', '', level, flags=re.IGNORECASE).strip().lower()
                            parts.append(f"{count} {clean_level} risk")
                    if parts:
                        total = sum(counts.values)
                        summary_line = ", ".join(parts) + " violation" + ("s" if total != 1 else "")
                        doc.add_paragraph(summary_line)

                # Add AI summaries as bullets, ordered high to low risk, with bold risk label
                if "ai_summary" in group.columns:
                    risk_order = {"High Risk": 0, "Moderate Risk": 1, "Low Risk": 2}

                    # Explode pipe-delimited comments and risk levels into individual rows
                    pairs = []
                    for _, vrow in group.iterrows():
                        raw_summary = str(vrow.get("comment", ""))
                        raw_risk = str(vrow.get("risk_level", ""))
                        if not raw_summary or raw_summary.lower() in ("nan", ""):
                            continue
                        summaries = [s.strip() for s in raw_summary.split(" | ") if s.strip()]
                        risks = [r.strip() for r in raw_risk.split(" | ") if r.strip()] if raw_risk and raw_risk.lower() not in ("nan", "") else []
                        for i, summary in enumerate(summaries):
                            risk = risks[i] if i < len(risks) else ""
                            sort_key = risk_order.get(risk.title(), 99)
                            pairs.append((sort_key, risk, summary))

                    # Sort all individual violations high to low
                    pairs.sort(key=lambda x: x[0])

                    for _, risk_label, summary in pairs:
                        p = doc.add_paragraph(style="List Bullet")
                        if risk_label and risk_label.upper() != "NA":
                            clean_label = re.sub(r'\s*risk\s*', '', risk_label, flags=re.IGNORECASE).strip().title()
                            p.add_run(f"{clean_label}: ").bold = True
                        p.add_run(summary)

        # In-compliance section
        heading_in = doc.add_paragraph()
        run_in = heading_in.add_run("Facilities with no violations:")
        run_in.bold = True
        run_in.font.size = Pt(16)

        if passed_deduped.empty:
            doc.add_paragraph("No passing inspections recorded.")
        else:
            for _, row in passed_deduped.iterrows():
                facility = str(row.get("facility", ""))
                address = str(row.get("address", ""))
                date = str(row.get("inspection_date", ""))
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(facility).bold = True
                p.add_run(f", {address}, on {date}")

        # Footer disclaimer
        footer_texts = [
            (
                "This post was automatically generated with ",
                "data",
                "http://cedatareporting.pa.gov/reports/powerbi/Public/AG/FS/PBI/Food_Safety_Inspections",
                " from the Pennsylvania Department of Agriculture\u2019s database of Food Safety Inspections for Retail Facilities. We have included violations and inspector comments only for facilities that were out of compliance in the previous week. We have also labeled violations as high, moderate, or low risk. These categories align directly with priority levels identified in the FDA Food Code: Priority, Priority Foundation, and Core."
            ),
        ]
        p = doc.add_paragraph()
        p.add_run("This post was automatically generated with ").italic = True
        add_hyperlink(p, "data", "http://cedatareporting.pa.gov/reports/powerbi/Public/AG/FS/PBI/Food_Safety_Inspections")
        p.add_run(" from the Pennsylvania Department of Agriculture\u2019s database of Food Safety Inspections for Retail Facilities. We have also labeled violations as high, moderate, or low risk. These categories align directly with priority levels identified in the FDA Food Code: Priority, Priority Foundation, and Core.").italic = True

        for footer_para in [
            "Priority items contribute directly to the elimination, prevention, or reduction to an acceptable level of hazards associated with foodborne illness or injury, such as handwashing, food handling, and temperature control, or other direct food contamination threats, such as rodents or pests. We noted violations of priority items as high risk.",
            "Priority foundation items support, facilitate, or enable control of risk factors that contribute to foodborne illness or injury, such as personnel training, labeling, and record-keeping. We noted violations of priority foundation items as moderate risk.",
            "Core items usually relate to standard operating procedures, facility structures, equipment design, or general maintenance. We noted violations of core items as low risk.",
        ]:
            p = doc.add_paragraph()
            p.add_run(footer_para).italic = True

        p = doc.add_paragraph()
        p.add_run("To read the full inspection reports, you can visit: ").italic = True
        add_hyperlink(p, "pafoodsafety.pa.gov/web/inspection/publicinspectionsearch.aspx", "https://www.pafoodsafety.pa.gov/web/inspection/publicinspectionsearch.aspx")

        # Save docx
        folder = os.path.join("data", "roundup", county_slug)
        os.makedirs(folder, exist_ok=True)
        output_path = os.path.join(folder, f"{date_slug}.docx")
        doc.save(output_path)
        print(f"Roundup saved: {output_path}")

        # Save plain text preview
        txt_path = output_path.replace(".docx", ".txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            for para in doc.paragraphs:
                f.write(para.text + "\n")
        print(f"Text preview saved: {txt_path}")

        # Upload to S3
        from helpers.uploader import upload_to_s3
        s3_key = f"2025/restaurant-inspections/roundup/{county_slug}/{date_slug}.docx"
        upload_to_s3(output_path, s3_key_override=s3_key)

        # Upload to Google Drive in county subfolder
        from helpers.gdrive_uploader import upload_to_gdrive, get_or_create_subfolder
        subfolder_id = get_or_create_subfolder(county_slug)
        gdrive_filename = f"{county_slug}-{date_slug}.docx"
        upload_to_gdrive(output_path, folder_id=subfolder_id, filename_override=gdrive_filename)

    except Exception as e:
        print(f"Roundup generation failed for {county_slug}: {e}")
        import traceback
        traceback.print_exc()