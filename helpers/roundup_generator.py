import os
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re as _re
from helpers.uploader import upload_to_s3


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


AP_MONTHS = {
    "January": "Jan.", "February": "Feb.", "March": "March",
    "April": "April", "May": "May", "June": "June",
    "July": "July", "August": "Aug.", "September": "Sept.",
    "October": "Oct.", "November": "Nov.", "December": "Dec."
}

AP_MONTHS_REVERSE = {v: k for k, v in AP_MONTHS.items()}

def reverse_ap_date(s):
    if not isinstance(s, str):
        return pd.NaT
    for abbr, full in AP_MONTHS_REVERSE.items():
        s = s.replace(abbr, full)
    return pd.to_datetime(s, errors="coerce")

def ap_date(dt, include_month=True):
    month = AP_MONTHS[dt.strftime("%B")]
    day = str(dt.day)
    return f"{month} {day}" if include_month else day

def get_week_range():
    today = datetime.today()
    start = today - timedelta(days=today.weekday() + 7)  # Monday of last week
    end = start + timedelta(days=6)  # Sunday of last week
    if start.month == end.month:
        date_range = f"{ap_date(start)}-{end.day}"
    else:
        date_range = f"{ap_date(start)}-{ap_date(end)}"
    return date_range, start.strftime("%Y-%m-%d"), start, end


def generate_roundup(file_path, county_slug):
    try:
        df = pd.read_excel(file_path)

        # Load inspections to get violation risk levels
        inspections_path = "data/inspections.xlsx"
        try:
            insp = pd.read_excel(inspections_path, dtype=str)
            insp["facility"] = insp["facility"].fillna("").str.strip()
            insp["city"] = insp["city"].fillna("").str.strip()
        except Exception as e:
            print(f"⚠️ Could not load inspections for violation counts: {e}")
            insp = pd.DataFrame()

        county_name = county_slug.title()
        date_range, date_slug, week_start, week_end = get_week_range()

        # Filter out-of-compliance to last completed week
        out = df[df["compliance"].str.strip().str.lower() == "out"] if "compliance" in df.columns else pd.DataFrame()
        if "last_inspection_date" in out.columns and not out.empty:
            out = out.copy()
            out["_parsed_date"] = out["last_inspection_date"].apply(reverse_ap_date)
            out = out[
                (out["_parsed_date"] >= pd.Timestamp(week_start.date())) &
                (out["_parsed_date"] <= pd.Timestamp(week_end.date()))
            ]
            out = out.drop(columns=["_parsed_date"])

        # Filter in-compliance to last completed week
        passed = df[df["compliance"].str.strip().str.lower() == "in"] if "compliance" in df.columns else df
        if "last_inspection_date" in passed.columns and not passed.empty:
            passed = passed.copy()
            passed["_parsed_date"] = passed["last_inspection_date"].apply(reverse_ap_date)
            passed = passed[
                (passed["_parsed_date"] >= pd.Timestamp(week_start.date())) &
                (passed["_parsed_date"] <= pd.Timestamp(week_end.date()))
            ]
            passed = passed.drop(columns=["_parsed_date"])

        # Build document
        doc = Document()

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title.add_run(f"{county_name} County Restaurant Inspections, {date_range}")
        run.bold = True
        run.font.size = Pt(20)

        # Date
        doc.add_paragraph(datetime.today().strftime("%B %-d, %Y"))

        # Intro
        intro = doc.add_paragraph(
            f"The Pennsylvania Department of Agriculture inspected the following food establishments "
            f"in {county_name} County this week. Find the full database at Spotlight PA's "
        )
        add_hyperlink(intro, "Restaurant Safety Tracker", "https://www.spotlightpa.org/restaurant-inspections")
        intro.add_run(".")

        # Out-of-compliance section
        heading_out = doc.add_paragraph()
        run_out = heading_out.add_run("Out-of-compliance inspections this week:")
        run_out.bold = True
        run_out.font.size = Pt(16)

        if out.empty:
            doc.add_paragraph("No out-of-compliance inspections this week.")
        else:
            for _, row in out.iterrows():
                facility = str(row.get("facility", ""))
                city = str(row.get("city", ""))
                address = str(row.get("address", ""))
                date = str(row.get("last_inspection_date", ""))

                facility_slug = _re.sub(r'[^a-z0-9]+', '-', facility.lower()).strip('-')
                tracker_url = f"http://localhost:8888/restaurant-inspections/#{facility_slug}"

                p = doc.add_paragraph()
                add_hyperlink(p, facility, tracker_url)
                addr_run = p.add_run(f"\n{address}")
                addr_run.italic = True
                date_run = p.add_run(f"\n{date}")
                date_run.italic = True

                # Count violations by risk level from inspections
                if not insp.empty:
                    insp_dates = insp["inspection_date"].apply(reverse_ap_date)
                    match = insp[
                        (insp["facility"] == facility) &
                        (insp["city"] == city) &
                        (insp_dates >= pd.Timestamp(week_start.date())) &
                        (insp_dates <= pd.Timestamp(week_end.date()))
                    ]
                    if not match.empty and "risk_level" in match.columns:
                        all_levels = (
                            match["risk_level"]
                            .dropna()
                            .astype(str)
                            .str.split(r"\s*\|\s*")
                            .explode()
                            .str.strip()
                            .str.title()
                        )
                        all_levels = all_levels[~all_levels.isin(["Na", "Nan", ""])]
                        counts = all_levels.value_counts()
                        if not counts.empty:
                            ap_nums = {
                                1: "One", 2: "Two", 3: "Three", 4: "Four",
                                5: "Five", 6: "Six", 7: "Seven", 8: "Eight", 9: "Nine"
                            }
                            for level, count in counts.items():
                                num = ap_nums.get(count, str(count))
                                clean_level = _re.sub(r'\s*risk\s*', '', level, flags=_re.IGNORECASE).strip().lower()
                                label = f"{clean_level} risk violation" + ("s" if count != 1 else "")
                                doc.add_paragraph(f"{num} {label}", style="List Bullet")

        # In-compliance section
        heading_in = doc.add_paragraph()
        run_in = heading_in.add_run("These establishments passed inspections this week:")
        run_in.bold = True
        run_in.font.size = Pt(16)

        if passed.empty:
            doc.add_paragraph("No passing inspections recorded this week.")
        else:
            for _, row in passed.iterrows():
                facility = str(row.get("facility", ""))
                address = str(row.get("address", ""))
                date = str(row.get("last_inspection_date", ""))
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(facility).bold = True
                p.add_run(f", {address}, {date}")

        # Save
        folder = os.path.join("data", "roundup", county_slug)
        os.makedirs(folder, exist_ok=True)
        output_path = os.path.join(folder, f"{date_slug}.docx")
        doc.save(output_path)
        print(f"Roundup saved: {output_path}")

        # Also save a plain text version for review
        txt_path = output_path.replace(".docx", ".txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            for para in doc.paragraphs:
                f.write(para.text + "\n")
        print(f"Text preview saved: {txt_path}")

        # Upload to S3
        s3_key_override = f"2025/restaurant-inspections/roundup/{county_slug}/{date_slug}.docx"
        upload_to_s3(output_path, s3_key_override=s3_key_override)

    except Exception as e:
        print(f"Roundup generation failed for {county_slug}: {e}")
        import traceback
        traceback.print_exc()